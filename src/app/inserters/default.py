from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL


def inserter(template, data):
    """a template for insertion of data into a docx template

    :param template: python-docx Document object made from a blank template
    :param data: information to be inserted into the template
    :return: nothing
    """

    # First page header table

    template.sections[0].first_page_header.tables[0].cell(0, 1).paragraphs[0].add_run(text=data.title)
    template.sections[0].first_page_header.tables[0].cell(0, 1).paragraphs[0].runs[0].bold = True
    template.sections[0].first_page_header.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(12)
    template.sections[0].first_page_header.tables[0].cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].add_run(text="Document No.:\n")
    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].runs[0].bold = True
    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(9)

    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].add_run(text=data.doc_no)
    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].runs[1].bold = False
    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].runs[1].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(1, 1).paragraphs[0].runs[1].font.size = Pt(11)
    template.sections[0].first_page_header.tables[0].cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].add_run(text="Prepared By:\n")
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[0].bold = True
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[0].font.size = Pt(9)

    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].add_run(text=f"{data.prep_by}\n")
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[1].bold = False
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[1].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[1].font.size = Pt(9)

    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].add_run(text=data.prep_by_title)
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[2].bold = False
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[2].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 0).paragraphs[0].runs[2].font.size = Pt(9)
    template.sections[0].first_page_header.tables[0].cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].add_run(text=f"Approved By:\n")
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[0].bold = True
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(9)

    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].add_run(text=f"{data.app_by}\n")
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[1].bold = False
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[1].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[1].font.size = Pt(9)

    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].add_run(text=data.app_by_title)
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[2].bold = False
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[2].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 1).paragraphs[0].runs[2].font.size = Pt(9)
    template.sections[0].first_page_header.tables[0].cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].add_run(text=f"Date Effective:\n")
    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].runs[0].bold = True
    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].runs[0].font.size = Pt(9)

    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].add_run(text=f"{data.date_eff}")
    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].runs[1].bold = False
    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].runs[1].font.name = "Arial"
    template.sections[0].first_page_header.tables[0].cell(2, 2).paragraphs[0].runs[1].font.size = Pt(11)
    template.sections[0].first_page_header.tables[0].cell(2, 2).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # following page header tables

    template.sections[0].header.tables[0].cell(0, 1).paragraphs[0].add_run(text=data.title)
    template.sections[0].header.tables[0].cell(0, 1).paragraphs[0].runs[0].bold = True
    template.sections[0].header.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].header.tables[0].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(12)
    template.sections[0].header.tables[0].cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].add_run(text="Document No.:\n")
    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].runs[0].bold = True
    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(9)

    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].add_run(text=data.doc_no)
    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].runs[1].bold = False
    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].runs[1].font.name = "Arial"
    template.sections[0].header.tables[0].cell(1, 1).paragraphs[0].runs[1].font.size = Pt(11)
    template.sections[0].header.tables[0].cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].add_run(text="Date Effective:\n")
    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].runs[0].bold = True
    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].runs[0].font.name = "Arial"
    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].runs[0].font.size = Pt(9)

    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].add_run(text=data.date_eff)
    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].runs[1].bold = False
    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].runs[1].font.name = "Arial"
    template.sections[0].header.tables[0].cell(1, 2).paragraphs[0].runs[1].font.size = Pt(11)
    template.sections[0].header.tables[0].cell(1, 2).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # equipment table

    template.tables[0].cell(0, 1).paragraphs[0].add_run(text=data.dept)
    template.tables[0].cell(0, 1).paragraphs[0].runs[0].bold = True
    template.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.tables[0].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(11)
    template.tables[0].cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    template.tables[0].cell(1, 1).paragraphs[0].add_run(text=data.equip_desc)
    template.tables[0].cell(1, 1).paragraphs[0].runs[0].bold = True
    template.tables[0].cell(1, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.tables[0].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(11)
    template.tables[0].cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    template.tables[0].cell(2, 1).paragraphs[0].add_run(text=data.equip_id)
    template.tables[0].cell(2, 1).paragraphs[0].runs[0].bold = True
    template.tables[0].cell(2, 1).paragraphs[0].runs[0].font.name = "Arial"
    template.tables[0].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(11)
    template.tables[0].cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # lockout table

    for i, row in enumerate(data.lockout, 1):

        try:
            test = template.tables[1].cell(i, 0).paragraphs[0].text
        except IndexError:
            template.tables[1].add_row()

        if row.context:
            other = template.tables[1].cell(i, 8)
            template.tables[1].cell(i, 0).merge(other)

            template.tables[1].cell(i, 0).paragraphs[0].add_run(text=row.context_text)
            template.tables[1].cell(i, 0).paragraphs[0].runs[0].bold = True
            template.tables[1].cell(i, 0).paragraphs[0].runs[0].font.name = "Arial"
            template.tables[1].cell(i, 0).paragraphs[0].runs[0].font.size = Pt(10)

        else:
            template.tables[1].cell(i, 0).paragraphs[0].add_run(text=row.num)
            template.tables[1].cell(i, 0).paragraphs[0].runs[0].font.name = "Arial"
            template.tables[1].cell(i, 0).paragraphs[0].runs[0].font.size = Pt(9)

            template.tables[1].cell(i, 1).paragraphs[0].add_run(text=row.point)
            template.tables[1].cell(i, 1).paragraphs[0].runs[0].font.name = "Arial"
            template.tables[1].cell(i, 1).paragraphs[0].runs[0].font.size = Pt(9)

            template.tables[1].cell(i, 2).paragraphs[0].add_run(text=row.tag_no)
            template.tables[1].cell(i, 2).paragraphs[0].runs[0].font.name = "Arial"
            template.tables[1].cell(i, 2).paragraphs[0].runs[0].font.size = Pt(9)

            template.tables[1].cell(i, 3).paragraphs[0].add_run(text=row.energy_src)
            template.tables[1].cell(i, 3).paragraphs[0].runs[0].font.name = "Arial"
            template.tables[1].cell(i, 3).paragraphs[0].runs[0].font.size = Pt(9)

            template.tables[1].cell(i, 6).paragraphs[0].add_run(text=row.isolating_means)
            template.tables[1].cell(i, 6).paragraphs[0].runs[0].font.name = "Arial"
            template.tables[1].cell(i, 6).paragraphs[0].runs[0].font.size = Pt(9)

    # isolation table (special precautions)

    template.tables[2].cell(4, 0).paragraphs[0].add_run(text="Special Precautions: ")
    template.tables[2].cell(4, 0).paragraphs[0].runs[0].bold = True
    template.tables[2].cell(4, 0).paragraphs[0].runs[0].font.name = "Arial"
    template.tables[2].cell(4, 0).paragraphs[0].runs[0].font.size = Pt(12)

    template.tables[2].cell(4, 0).paragraphs[0].add_run(text=data.special_precautions)
    template.tables[2].cell(4, 0).paragraphs[0].runs[1].bold = False
    template.tables[2].cell(4, 0).paragraphs[0].runs[1].font.name = "Arial"
    template.tables[2].cell(4, 0).paragraphs[0].runs[1].font.size = Pt(12)

    # extra steps added to section 3.0 Requirements

    if len(data.extra) > 0:
        for j, line in enumerate(data.extra):
            if j == 0:
                template.paragraphs[4].add_run(text="\n")
                template.paragraphs[4].runs[4].bold = False
                template.paragraphs[4].runs[4].font.name = "Arial"
                template.paragraphs[4].runs[4].font.size = Pt(10)
                continue
            template.paragraphs[4].add_run(text=f"3.{j}\t{data.extra[j-1]}\n")
            template.paragraphs[4].runs[j + 4].bold = False
            template.paragraphs[4].runs[j + 4].font.name = "Arial"
            template.paragraphs[4].runs[j + 4].font.size = Pt(10)
